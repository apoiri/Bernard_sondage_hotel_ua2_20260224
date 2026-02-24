# -*- coding: utf-8 -*-
"""
Module 11 – Synthèse Word à l'intention du professeur
Simulation « L'Hôtel Boutique Art de Vivre »
Génère un document Word listant tous les éléments d'apprentissage que les étudiants
peuvent réaliser avec le fichier sondage_hotel_data.csv, pour que le professeur
puisse couvrir toutes les activités pédagogiques.
"""

from pathlib import Path

# ========== HYPOTHÈSES DU SCRIPT (modifiables par l'opérateur) ==========
FICHIER_WORD_SORTIE = "synthese_elements_apprentissage_prof.docx"
REPERTOIRE = None   # None = même dossier que ce script
FICHIER_CSV_REFERENCE = "sondage_hotel_data.csv"

# ========== NE PAS MODIFIER CI-DESSOUS (logique du script) ==========

import sys
sys.path.insert(0, str(Path(__file__).resolve().parent))

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.shared import RGBColor
except ImportError:
    raise ImportError("Module python-docx requis. Installer avec : pip install python-docx")

# Couleur du texte pour garantir la visibilité (éviter texte invisible en thème sombre)
COULEUR_TEXTE = RGBColor(0, 0, 0)


def _repertoire():
    return Path(REPERTOIRE) if REPERTOIRE else Path(__file__).resolve().parent


def _add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = COULEUR_TEXTE
    return p


def _add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.color.rgb = COULEUR_TEXTE
    if bold:
        run.bold = True
    return p


def run():
    """Génère le document Word de synthèse pour le professeur."""
    doc = Document()
    rep = _repertoire()
    chemin_sortie = rep / FICHIER_WORD_SORTIE
    # Forcer le style Normal : police noire, taille 11 pt (visibilité dans tous les lecteurs)
    try:
        style = doc.styles["Normal"]
        style.font.color.rgb = COULEUR_TEXTE
        style.font.size = Pt(11)
    except Exception:
        pass

    # ---- Titre ----
    _add_heading(doc, "Synthèse – Éléments d'apprentissage", level=0)
    _add_para(doc, "Document à l'intention du professeur – Simulation « L'Hôtel Boutique Art de Vivre »", bold=True)
    _add_para(doc, "Ce document recense tous les éléments d'apprentissage que les étudiants peuvent réaliser avec le fichier de données généré, afin d'assurer la couverture complète des objectifs pédagogiques en analyse de sondage et prise de décision.")
    doc.add_paragraph()

    # ---- 1. Synthèse du rapport de validation des analyses (en premier pour visibilité) ----
    _add_heading(doc, "1. Synthèse du rapport de validation des analyses", level=1)
    _add_para(doc, "Le rapport complet (RAPPORT_DETAILLE_VALIDATION_ANALYSES.md) est fourni avec ce livrable. Les chiffres ci-dessous ont été vérifiés sur le fichier sondage_hotel_data.csv (script verif_calculs_stats.py et module m10) : ils correspondent à l'analyse du fichier livré, et non à une simple reprise de modèle.", bold=False)
    _add_para(doc, "Résultats attendus et lien avec la décision :", bold=True)
    synthèse = [
        "Proportions et IC 95 % : % annulées 18,7 %, Rev_Spa>0 52,3 %, Forfait Gastro 42,1 %. IC resserrés (n ≈ 10 500). → Objectifs réalistes, suivre les écarts.",
        "Khi-deux (Segment × Type_Forfait) : χ² ≈ 2842, p ≈ 0. Liaison très forte (Loisirs → Gastronomique, Congressiste → Chambre seule). → Cibler l'offre par segment.",
        "Pearson (Rev_Spa × Satisfaction_NPS) : en brut r ≈ 0,45 ; après nettoyage r ≈ 0,76. Dire aux étudiants « r ≈ 0,75 après nettoyage ». → Leviers satisfaction / fidélisation.",
        "ANOVA (Rev_Resto par Segment) : F ≈ 10 363, p ≈ 0. Différences très significatives entre segments. → Allouer ressources restaurant par segment.",
        "Régression (Total_Facture ~ Rev_Chambre + Rev_Resto + Rev_Spa) : coefficients ≈ 1,00 ; 1,08 ; 1,29, R² ≈ 0,9997. → Piloter structure des revenus.",
        "Annulations par canal : Direct ~6,8 %, Intermédiaire ~30,5 %. → Risque OTA, développer le direct.",
        "Anomalies pédagogiques : 15 Externes avec Nuits>0 ; 4 Annulee=Oui avec Nuits>0 ; 3 NPS=99 ; 5 % manquants Rev_Spa/Rev_Resto ; 10 doublons. → Nettoyage obligatoire avant analyse.",
        "Tarification active : utiliser Saison_calendrier (pas Saison). Haute ~412 $, Basse ~265 $ (+55 %). Détail : Note_Saison_vs_Tarification_Dynamique.docx.",
    ]
    for s in synthèse:
        _add_para(doc, f"  • {s}")
    doc.add_paragraph()

    # ---- 2. Présentation du jeu de données ----
    _add_heading(doc, "2. Présentation du jeu de données", level=1)
    _add_para(doc, f"Fichier de travail : {FICHIER_CSV_REFERENCE}")
    try:
        from m01_config import get_config
        _c = get_config()
        _taux_pct = _c.get("TAUX_OCCUPATION_POURCENT", 72)
        _taux_industrie = _c.get("TAUX_INDUSTRIE_POURCENT", 73)
    except Exception:
        _taux_pct = 72
        _taux_industrie = 73
    _add_para(doc, f"Contenu : Réservations simulées sur une année complète pour un hôtel de 100 chambres. Le taux d'occupation cible est de {_taux_pct} % (hypothèse : performance de l'hôtel = industrie {_taux_industrie} % ± 1 %). Ce taux est le paramètre de la simulation pour fixer le volume de réservations ; le taux calculable à partir du CSV (nuits vendues / capacité 100×365) peut différer (annulations, Externes, périmètre). Le fichier contient environ 10 500 lignes (réservations) plus des lignes dupliquées volontaires pour l'exercice de nettoyage.")
    _add_para(doc, "Variables principales :")
    variables = [
        ("ID_Client", "Identifiant unique (ex. AB-00001)"),
        ("Segment", "Persona : Affaires_Solo, Congressiste, Loisirs_Couple, Local_Gourmet, Local_Spa"),
        ("Type_Client", "Interne (séjourne) / Externe (ne séjourne pas)"),
        ("Saison", "Haute / Basse (selon segment ; ne pas utiliser pour les analyses de tarification dynamique)."),
        ("Saison_calendrier", "Basse / Épaule / Haute (dérivée du mois ; à utiliser pour les exercices sur la tarification active et les prix)."),
        ("Canal_reservation", "Source : Site Web, Courriel, Téléphone, Booking.com, Expedia, etc."),
        ("Type_canal", "Direct / Intermediaire"),
        ("Annulee", "Oui / Non (taux annulation ~18 %, plus élevé via intermédiaires)"),
        ("Nuits", "Nombre de nuits (0 si Externe ou annulé)"),
        ("Rev_Chambre", "Revenus chambre ($)"),
        ("Rev_Banquet", "Revenus banquet ($)"),
        ("Rev_Resto", "Revenus restaurant ($)"),
        ("Rev_Spa", "Revenus spa ($)"),
        ("Type_Forfait", "Forfait Gastronomique / Chambre Seule"),
        ("Satisfaction_NPS", "Score NPS 0–10 (NaN si annulé)"),
        ("Total_Facture", "Montant total ($)"),
    ]
    for nom, desc in variables:
        _add_para(doc, f"  • {nom} : {desc}")
    doc.add_paragraph()

    # ---- 3. Éléments d'apprentissage (tableau) ----
    _add_heading(doc, "3. Éléments d'apprentissage à réaliser avec le fichier", level=1)
    _add_para(doc, "Le tableau ci-dessous liste chaque technique ou objectif pédagogique, les variables à utiliser et le type de sortie attendu. Le professeur peut s'en servir pour construire les exercices ou vérifier que tous les éléments sont couverts.")
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Élément d'apprentissage"
    hdr[1].text = "Variables utilisées"
    hdr[2].text = "Objectif / Sortie attendue"
    hdr[3].text = "Coché"
    for cell in hdr:
        for r in cell.paragraphs[0].runs:
            r.bold = True
            r.font.color.rgb = COULEUR_TEXTE

    lignes = [
        ("Statistiques descriptives", "Toutes (effectifs, moyennes, écarts-types, min/max)", "Tableaux de synthèse par segment, Type_Client, Saison, Saison_calendrier, Annulee ; description des variables numériques (Nuits, Rev_*, Total_Facture, Satisfaction_NPS).", "☐"),
        ("Proportions et intervalles de confiance", "Annulee ; Rev_Spa ; Type_Forfait", "Estimer une proportion en population (ex. % annulées, % utilisant le spa, % Forfait Gastronomique) et calculer l'IC 95 %.", "☐"),
        ("Test du Khi-deux (indépendance)", "Segment × Type_Forfait", "Vérifier la liaison entre segment et type de forfait (Loisirs → Forfait Gastronomique, Congressiste → Chambre Seule) ; khi², ddl, p-value.", "☐"),
        ("Corrélation de Pearson", "Rev_Spa, Satisfaction_NPS (lignes non annulées)", "Mesurer la liaison linéaire entre dépense spa et satisfaction ; r, p-value. Attendu : corrélation positive (données conçues pour r ≈ 0,75 après nettoyage ; en brut r ≈ 0,45).", "☐"),
        ("Tarification active (prix selon la période)", "Saison_calendrier, Tarif_applique ou Rev_Chambre (Interne, Non annulée, Nuits > 0)", "Comparer les tarifs ou revenus chambre selon Saison_calendrier (Basse / Épaule / Haute). Attendu : Haute > Épaule > Basse (~+55 % Haute vs Basse). Ne pas utiliser la colonne Saison pour cet exercice.", "☐"),
        ("Comparaison de moyennes (ANOVA)", "Rev_Resto par Segment", "Comparer les dépenses restaurant entre segments ; F, p-value. Attendu : différences significatives entre segments (p < 0,05).", "☐"),
        ("Régression linéaire multiple", "Total_Facture ~ Rev_Chambre + Rev_Resto + Rev_Spa", "Prédire la facture totale à partir des revenus par poste ; coefficients, R², interprétation. Attendu : coefficients proches de 1 ; 1,1 ; 1,3.", "☐"),
        ("Annulations par canal", "Type_canal, Annulee ; Canal_reservation", "Comparer les taux d'annulation Direct vs Intermédiaire ; tableaux et graphiques. Attendu : ~7 % direct, ~29 % intermédiaire.", "☐"),
        ("Nettoyage des données (data cleaning)", "Type_Client & Nuits ; Annulee & Nuits ; Satisfaction_NPS ; Rev_Spa, Rev_Resto ; doublons", "Détecter et traiter : Externes avec Nuits > 0 (incohérence) ; Annulee = Oui avec Nuits > 0 (4 lignes) ; valeurs aberrantes (ex. NPS = 99) ; valeurs manquantes (Rev_Spa, Rev_Resto) ; lignes dupliquées.", "☐"),
        ("Prise de décision / recommandations", "Segments, revenus, satisfaction, canal", "Synthétiser les résultats pour formuler des recommandations (ex. cibler les canaux à faible annulation, segments à forte dépense spa, etc.).", "☐"),
    ]
    for elem, vars_, objectif, coche in lignes:
        row = table.add_row()
        for i, val in enumerate((elem, vars_, objectif, coche)):
            cell = row.cells[i]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = COULEUR_TEXTE
    doc.add_paragraph()

    # ---- 4. Anomalies pédagogiques ----
    _add_heading(doc, "4. Anomalies pédagogiques (pour l'exercice de nettoyage)", level=1)
    _add_para(doc, "Le fichier contient volontairement les anomalies suivantes, à faire détecter et corriger par les étudiants :")
    anomalies = [
        "Incohérence logique : des clients « Externes » ont Nuits > 0 (une quinzaine de lignes).",
        "Incohérence logique : Annulee = Oui avec Nuits > 0 (4 lignes).",
        "Outliers : quelques lignes avec Satisfaction_NPS = 99 (erreur de saisie simulée).",
        "Valeurs manquantes : environ 5 % de cellules vides dans Rev_Spa et Rev_Resto.",
        "Doublons : une dizaine de lignes sont dupliquées (à identifier et supprimer ou fusionner).",
    ]
    for a in anomalies:
        _add_para(doc, f"  • {a}")
    _add_para(doc, "Après nettoyage, les analyses (proportions, régression, Pearson, etc.) peuvent être refaites pour comparer les résultats avec/sans anomalies.")
    doc.add_paragraph()

    # ---- 5. Checklist professeur ----
    _add_heading(doc, "5. Checklist – Éléments traités en cours / TD", level=1)
    _add_para(doc, "Le professeur peut cocher (manuellement dans ce document) les éléments d'apprentissage déjà traités avec les étudiants, afin de s'assurer que tous les objectifs du fichier sont couverts.")
    _add_para(doc, "Les lignes du tableau de la section 3 contiennent une colonne « Coché » à cocher au fur et à mesure.")
    doc.add_paragraph()
    _add_para(doc, "— Fin du document —", bold=True)

    doc.save(chemin_sortie)
    return chemin_sortie


if __name__ == "__main__":
    print("=" * 60)
    print("HYPOTHÈSES MODULE 11 (Synthèse Word professeur)")
    print("=" * 60)
    print(f"  Fichier Word sortie = {FICHIER_WORD_SORTIE}")
    print("=" * 60)
    chemin = run()
    print(f"Document Word enregistré : {chemin}")

# ========== RÈGLE DE BLOCAGE ==========
# Aucune modification de ce module n'est acceptée sans l'autorisation du maître d'ouvrage du projet.
