# -*- coding: utf-8 -*-
"""
Génère le document Word équivalent à NOTE_SAISON_VS_TARIFICATION_DYNAMIQUE.md
pour communication au client (professeur). Le fichier .md reste la référence projet.
"""
from pathlib import Path

REP = Path(__file__).resolve().parent
FICHIER_SORTIE = REP / "Note_Saison_vs_Tarification_Dynamique.docx"

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    raise ImportError("python-docx requis : pip install python-docx")


def _para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p


def run():
    doc = Document()
    doc.add_heading("Saison vs tarification dynamique – Note pour vérification et pédagogie", level=0)
    _para(doc, "Cette note répond au diagnostic d'une vérification indépendante de la tarification active et précise comment interpréter les colonnes Saison et Saison_calendrier dans le CSV.")
    doc.add_paragraph()

    doc.add_heading("1. Ce que la vérification indépendante a constaté", level=1)
    _para(doc, "Tarification active (dynamique) : OUI — elle est très nette quand on regarde le mois du séjour (Mois_sejour) : été (juin–août) nettement plus cher que hiver (déc–fév), +50 à 55 % selon les types de chambre.")
    _para(doc, "Colonne Saison (Haute/Basse) : elle ne reflète pas cette hausse. En comparant Tarif_applique selon Saison, Haute et Basse sont quasi identiques (non significatif). Donc, avec la seule variable Saison, on ne peut pas conclure à une tarification active.")
    _para(doc, "Raison technique :", bold=True)
    _para(doc, "• Saison vient du module 2 (segments) : elle est tirée selon le profil client (ex. Congressiste → plus souvent « Basse », Loisirs → plus souvent « Haute »), pas selon le calendrier.")
    _para(doc, "• Tarif_applique et Rev_Chambre sont calculés par le module de tarification dynamique à partir du mois (Mois_sejour) et des coefficients par mois, sans utiliser la colonne Saison.")
    _para(doc, "D'où le décalage : Saison = saison « marketing » par segment ; tarification dynamique = pilotée par le mois.")
    doc.add_paragraph()

    doc.add_heading("2. Modifications apportées pour aligner analyse et pédagogie", level=1)
    doc.add_heading("Nouvelle colonne : Saison_calendrier", level=2)
    _para(doc, "Une colonne Saison_calendrier a été ajoutée au CSV, dérivée du mois (Mois_sejour), pour que les analyses « haute vs basse saison » soient cohérentes avec la tarification :")
    table1 = doc.add_table(rows=4, cols=2)
    table1.style = "Table Grid"
    h = table1.rows[0].cells
    h[0].text = "Saison_calendrier"
    h[1].text = "Mois (Mois_sejour)"
    table1.rows[1].cells[0].text = "Haute"
    table1.rows[1].cells[1].text = "6 (juin), 7 (juillet), 8 (août)"
    table1.rows[2].cells[0].text = "Basse"
    table1.rows[2].cells[1].text = "12 (décembre), 1 (janvier), 2 (février)"
    table1.rows[3].cells[0].text = "Épaule"
    table1.rows[3].cells[1].text = "3, 4, 5, 9, 10, 11 (mars–mai, sept.–nov.)"
    doc.add_paragraph()
    _para(doc, "Avec Saison_calendrier :")
    _para(doc, "• Les comparaisons « Haute vs Basse » sur Tarif_applique (ou Rev_Chambre) montrent bien la tarification active (été plus cher que hiver).")
    _para(doc, "• L'effet reste visible par type de chambre (environ +50–55 % en Haute vs Basse).")
    doc.add_heading("Colonne Saison (inchangée)", level=2)
    _para(doc, "La colonne Saison (Haute/Basse) est conservée telle quelle : elle reste basée sur le segment client et continue d'alimenter d'autres modules (ex. revenus centres de profit, forfait). Elle ne doit pas être utilisée pour analyser la tarification dynamique ; pour cela, utiliser Saison_calendrier ou Mois_sejour.")
    doc.add_paragraph()

    doc.add_heading("3. Recommandations pour les analyses et la formation", level=1)
    _para(doc, "Pour analyser la tarification active (prix selon la période) : utiliser Mois_sejour ou Saison_calendrier (et non Saison). Exemples : Tarif_applique moyen par Saison_calendrier (Haute > Basse) ; comparaison été (6–8) vs hiver (12–2) sur Tarif_applique ou Rev_Chambre.")
    _para(doc, "Pour les exercices pédagogiques : Saison = liaison avec le segment, comportement client, autres indicateurs (hors tarification dynamique). Saison_calendrier (ou Mois_sejour) = tarification dynamique, effet « haute vs basse » sur les prix et revenus chambre.")
    _para(doc, "Pour Ottawa / Québec : le codage retenu (Haute = 6–8, Basse = 12–2, Épaule = reste) peut être adapté dans le module m02e_saison_calendrier.py si vous fixez d'autres mois officiels pour haute/basse saison.")
    doc.add_paragraph()

    doc.add_heading("4. Résumé pour le client et les professeurs", level=1)
    table2 = doc.add_table(rows=4, cols=2)
    table2.style = "Table Grid"
    table2.rows[0].cells[0].text = "Élément"
    table2.rows[0].cells[1].text = "Conclusion"
    table2.rows[1].cells[0].text = "Tarification active dans le fichier"
    table2.rows[1].cells[1].text = "Oui — très nette via Mois_sejour (été vs hiver, +50–55 % par type de chambre)."
    table2.rows[2].cells[0].text = "Colonne Saison (Haute/Basse)"
    table2.rows[2].cells[1].text = "Reflète la saison par segment, pas le calendrier ; ne pas l'utiliser pour la tarification dynamique."
    table2.rows[3].cells[0].text = "Colonne Saison_calendrier"
    table2.rows[3].cells[1].text = "Dérivée du mois ; alignée sur la logique de tarification ; à utiliser pour analyser « haute vs basse » sur les tarifs et revenus chambre."
    doc.add_paragraph()
    _para(doc, "Les données sont valides pour démontrer la tarification active ; l'usage de Saison_calendrier (ou Mois_sejour) permet de le faire de façon claire et cohérente pour les étudiants et les vérifications indépendantes.")

    doc.save(FICHIER_SORTIE)
    print(f"Document Word enregistré : {FICHIER_SORTIE}")
    return str(FICHIER_SORTIE)


if __name__ == "__main__":
    run()
